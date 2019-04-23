import datetime
import sys
import time

from PyQt5.QtCore import *
from docx import Document
from qtpy.QtWidgets import *
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.by import By

from ui.AccDialog import Ui_AccDialog
from ui.Dialog import Ui_Dialog
from ui.mainwindow import Ui_MainWindow

app = QApplication(sys.argv)


class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)

        self.accounts = {}
        with open('accountLists.txt', 'r') as f:
            d = []
            for i in f:
                if '=' in i:
                    d = i.strip().split(' ')
                    self.accounts[d[0]] = []
                elif '=' not in i and not len(i.strip()) == 0:
                    i = i.strip()
                    self.accounts[d[0]].append(i.strip())
                else:
                    continue

        self.currentBrowser = ''
        self.currentPassword = ''
        self.newPassword = ''
        self.finish_dialog = QMessageBox()
        self.error_dialog = QMessageBox()
        self.docx_dialog = QMessageBox()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.addCombo()
        browsers = ['Chrome', 'Firefox']
        self.ui.comboBox_2.addItems(browsers)
        self.ui.label_5.hide()
        self.ui.label_4.hide()
        self.ui.pushButton.clicked.connect(self.changePW)
        self.ui.addButton.clicked.connect(self.addList)
        self.ui.showButton.clicked.connect(self.showAcc)

    def addCombo(self):
        for key, value in self.accounts.items():
            self.ui.comboBox.addItem(key)

    def changePW(self):
        self.currentBrowser = self.ui.comboBox_2.currentText()
        self.currentPassword = self.ui.inputCurrent.text()
        self.newPassword = self.ui.inputNew.text()
        self.account = self.ui.comboBox.currentText()
        self.obj = Worker(self.account, self.currentBrowser, self.currentPassword, self.newPassword, self.accounts)
        self.obj.message.connect(self.errorDialog)
        self.obj.finished.connect(self.done)
        self.obj.progress.connect(self.loadingBar)
        self.obj.docFinish.connect(self.docxDialog)
        self.obj.label4.connect(self.okay4)
        self.obj.label5.connect(self.okay5)
        self.obj.start()

    def addList(self):
        self.Adddialog = QDialog()
        self.uiAddDialog = Ui_Dialog()
        self.uiAddDialog.setupUi(self.Adddialog)
        self.Adddialog.show()
        self.uiAddDialog.addRow.clicked.connect(self.addCell)
        rsp = self.Adddialog.exec_()
        if rsp == QDialog.Accepted:
            listname = self.uiAddDialog.listName.text()
            if listname:
                if not listname in self.accounts:
                    self.accounts[listname] = []
                    with open('accountLists.txt', 'a') as f:
                        f.write('\n' + listname + ' ' + '=\n')

                    print(self.accounts)
            for i in range(0, self.uiAddDialog.accTable.rowCount()):
                if not self.uiAddDialog.accTable.item(i, 0) == None:
                    s = self.uiAddDialog.accTable.item(i, 0).text()
                    with open('accountLists.txt', 'a') as f:
                        f.write(s + '\n')


                else:
                    print('leer')


        else:
            print('Cancel')

    def addCell(self):
        row = self.uiAddDialog.accTable.rowCount()
        self.uiAddDialog.accTable.insertRow(row)

    def showAcc(self):
        self.Accdialog = QDialog()
        self.uiAccDialog = Ui_AccDialog()
        self.uiAccDialog.setupUi(self.Accdialog)
        for key, value in self.accounts.items():
            self.uiAccDialog.comboBox.addItem(key)
        if self.uiAccDialog.comboBox.currentText() in self.accounts:
            self.on_comboBox_changed(self.uiAccDialog.comboBox.currentText())
        self.uiAccDialog.comboBox.currentTextChanged.connect(self.on_comboBox_changed)
        self.Accdialog.show()
        self.Accdialog.exec_()

    def on_comboBox_changed(self, key):
        while self.uiAccDialog.tableWidget.rowCount() > 0:
            self.uiAccDialog.tableWidget.removeRow(0)
        if self.uiAccDialog.comboBox.currentText() == key:
            for i, j in self.accounts.items():
                if key == i:
                    for s in j:
                        row = self.uiAccDialog.tableWidget.rowCount()
                        self.uiAccDialog.tableWidget.insertRow(row)
                        self.uiAccDialog.tableWidget.setItem(row, 0, QTableWidgetItem(str(s)))

    def errorDialog(self, errorText):
        print(errorText)
        self.error_dialog.setIcon(QMessageBox.Critical)
        self.error_dialog.setText('Problem with password in account:' + errorText)
        self.error_dialog.show()

    def done(self, str):
        self.finish_dialog.setIcon(QMessageBox.Information)
        self.finish_dialog.setText('Passwords changed to:   ' + str)
        self.finish_dialog.show()

    def loadingBar(self, percentage):
        print(percentage)
        self.ui.progressBar.setValue(percentage)

    def docxDialog(self, str):
        self.docx_dialog.setIcon(QMessageBox.Information)
        self.docx_dialog.setText(str)
        self.docx_dialog.show()

    def okay4(self, okay1):
        self.ui.label_4.setText(okay1)
        self.ui.label_4.show()

    def okay5(self, okay2):
        self.ui.label_5.setText(okay2)
        self.ui.label_5.show()


class Worker(QThread):
    message = pyqtSignal(str)
    finished = pyqtSignal(str)
    progress = pyqtSignal(float)
    label4 = pyqtSignal(str)
    label5 = pyqtSignal(str)
    docFinish = pyqtSignal(str)

    def __init__(self, account, brwser, curPass, newPass, accountList):
        QThread.__init__(self)

        self.currentBrowser = brwser
        self.currentPassword = curPass
        self.newPassword = newPass
        self.account = account
        self.accountList = accountList

    def run(self):
        document = Document()

        table = document.add_table(rows=1, cols=2)
        document.add_paragraph(str(datetime.datetime.now()))
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'User Account'
        hdr_cells[1].text = 'Password (also for IoT Ext. and Integration)'

        count = float(0)
        self.progress.emit(count)
        while count < 100:
            for acc, address in self.accountList.items():
                if self.account == acc:
                    print(acc)
                    for user in self.accountList[acc]:
                        print(acc)
                        b = len(self.accountList[acc])
                        print(b)
                        count += 100 / b
                        page = "https://www2.industrysoftware.automation.siemens.com/webkey/"
                        if self.currentBrowser == 'Chrome':
                            options = webdriver.ChromeOptions()
                            driver = webdriver.Chrome(executable_path='webdriver/windows/chromedriver',
                                                      options=options)

                        if self.currentBrowser == 'Firefox':
                            options = webdriver.FirefoxOptions()
                            driver = webdriver.Firefox(executable_path='webdriver/windows/geckodriver',
                                                       options=options)
                        # #options.add_argument('-headless')  # run in background
                        driver.implicitly_wait(5)
                        # # driver.minimize_window()
                        driver.get(page)
                        driver.find_element(By.XPATH, "//tr[5]/td[2]/ul/font/li/a/font").click()
                        driver.find_element(By.XPATH, "//td[2]/input").click()
                        driver.find_element(By.NAME, "WebKey_Username").send_keys(user)
                        driver.find_element(By.NAME, "Existing_WebKey_Password").send_keys(self.currentPassword)
                        driver.find_element(By.NAME, "pass").click()
                        driver.find_element(By.NAME, "pass").send_keys(self.newPassword)
                        driver.find_element(By.NAME, "repass").click()
                        driver.find_element(By.NAME, "repass").send_keys(self.newPassword)
                        driver.find_element(By.XPATH, "//div[3]/div[2]/div/form/fieldset/input").click()
                        time.sleep(3)

                        # if wrong password entered
                        try:
                            driver.find_element(By.XPATH, "//h2[contains(.,'WebKey Error')]")
                            self.message.emit(user)
                            self.progress.emit(count)
                            row_cells = table.add_row().cells
                            row_cells[0].text = user
                            row_cells[1].text = 'Wrong Password entered!'
                            empty_cells = table.add_row().cells
                            empty_cells[0].text = ''
                            empty_cells[1].text = ''
                            driver.quit()
                            continue
                        except NoSuchElementException:
                            pass

                        # if new password matches older one

                        try:
                            driver.find_element(By.XPATH,
                                                "//h2[contains(.,'The following message was returned from the WebKey Server:')]")
                            self.message.emit(user)
                            self.progress.emit(count)
                            row_cells = table.add_row().cells
                            row_cells[0].text = user
                            row_cells[1].text = 'New Password matches older one!'
                            empty_cells = table.add_row().cells
                            empty_cells[0].text = ''
                            empty_cells[1].text = ''
                            driver.quit()
                            continue
                        except NoSuchElementException:
                            pass

                        # if driver.find_element(By.XPATH, ): Your Password has been Changed
                        try:
                            driver.find_element(By.XPATH, "//h2[contains(.,'Your Password has been Changed')]")
                            self.label4.emit('OK')
                            self.label5.emit('OK')
                            driver.quit()
                            row_cells = table.add_row().cells
                            row_cells[0].text = user
                            row_cells[1].text = self.newPassword
                            empty_cells = table.add_row().cells
                            empty_cells[0].text = ''
                            empty_cells[1].text = ''
                            self.progress.emit(count)
                            print(str(count) + '%')
                        except NoSuchElementException:
                            pass
        self.finished.emit(self.newPassword)
        for acc in self.accountList:
            if acc == self.account:
                document.save(acc + '.docx')
        self.docFinish.emit('Word file with all changed accounts created and saved in application folder! :)')


window = MainWindow()
window.show()

sys.exit(app.exec_())
